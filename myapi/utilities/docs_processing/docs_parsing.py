import pdfplumber
import docx

def extract_pdf_content(file_path):
    text_content = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()

                if page_text:
                    text_content.append(page_text)

        return "\n".join(text_content) if text_content else None

    except Exception as e:
        print(f"Error parsing PDF at {file_path}: {e}")
        return None


def extract_docx_content(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = []

        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text)

        for t in doc.tables:
            for r in t.rows:
                row_data= [cell.text.strip() for cell in r.cells if cell.text.strip()]
                if row_data:
                    full_text.append(" | ".join(row_data))
        return "\n".join(full_text) if full_text else None
    
    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        return None
    

def extract_text_content(file_path):
    try:
        with open(file_path, 'r', encoding= 'utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, 'r', encoding= 'latin-1') as f:
            return f.read()



def extract_content(file_path, file_type):
    if file_type.lower()== "pdf":
        return extract_pdf_content(file_path)
    elif file_type.lower() in ['docx', 'doc']:
        return extract_docx_content(file_path)
    elif file_type.lower()== 'txt':
        return extract_text_content(file_path)
    return None